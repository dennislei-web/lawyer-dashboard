"""將 1-on-1 工作表的 markdown 轉成 Word .docx
- 保留標題、粗體、表格、checkbox、分隔線
- Traditional Chinese font: Microsoft JhengHei
- A4 頁面，適合列印手寫
"""
import re
import sys
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")


def set_cell_shading(cell, color_hex):
    """設定 cell 背景色"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def apply_font(run, font_name="Microsoft JhengHei", size_pt=None, bold=None, color=None):
    """套用字體、字級、粗體、顏色"""
    if font_name:
        run.font.name = font_name
        # 中文字型需要 eastAsia 設定
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_inline(paragraph, text, base_size=11, default_font="Microsoft JhengHei"):
    """解析單行 inline 格式：**粗體** 和 `code`，寫成多個 run"""
    # Tokenize by bold markers and backticks
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            apply_font(run, default_font, base_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            apply_font(run, "Consolas", base_size - 1)
            run.font.color.rgb = RGBColor(0x88, 0x44, 0x00)
        else:
            run = paragraph.add_run(part)
            apply_font(run, default_font, base_size)


def add_horizontal_rule(doc):
    """在文件加一條水平線"""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "999999")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def parse_table_block(lines):
    """從 markdown 表格行提取 rows（list of list of str）"""
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    # 去掉分隔行（|---|---|）
    rows = [r for r in rows if not all(re.match(r"^:?-+:?$", c) for c in r if c)]
    return rows


def add_table_from_md(doc, md_rows):
    """建立 Word 表格，第一行為 header"""
    if not md_rows:
        return
    n_cols = len(md_rows[0])
    tbl = doc.add_table(rows=len(md_rows), cols=n_cols)
    tbl.style = "Table Grid"
    for i, row_cells in enumerate(md_rows):
        row = tbl.rows[i]
        for j, cell_text in enumerate(row_cells):
            if j >= n_cols:
                continue
            cell = row.cells[j]
            # 清空預設段落
            cell.text = ""
            p = cell.paragraphs[0]
            if i == 0:
                # Header — 粗體 + 淺藍底色
                set_cell_shading(cell, "DEEBF7")
                add_inline(p, f"**{cell_text}**" if not cell_text.startswith("**") else cell_text, base_size=10)
            else:
                add_inline(p, cell_text, base_size=10)


def md_to_docx(md_path, docx_path):
    src = Path(md_path).read_text(encoding="utf-8")
    lines = src.split("\n")

    doc = Document()

    # === 頁面設定：A4，適合列印 ===
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # 預設字型
    style = doc.styles["Normal"]
    style.font.name = "Microsoft JhengHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 空行 → 段落空行
        if not stripped:
            i += 1
            continue

        # 水平線
        if re.match(r"^-{3,}$", stripped) or re.match(r"^={3,}$", stripped):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            p = doc.add_paragraph()
            # 移除標題裡面的 ** 直接粗體
            clean_text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            run = p.add_run(clean_text)
            size = {1: 20, 2: 16, 3: 13}.get(level, 11)
            apply_font(run, size_pt=size, bold=True)
            if level == 1:
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            elif level == 2:
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            elif level == 3:
                run.font.color.rgb = RGBColor(0x30, 0x58, 0xA0)
            p.paragraph_format.space_before = Pt(8 if level <= 2 else 4)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # 表格（連續 | ... | 行）
        if stripped.startswith("|") and stripped.endswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                block.append(lines[i])
                i += 1
            rows = parse_table_block(block)
            add_table_from_md(doc, rows)
            # 表格後加個小間距
            doc.add_paragraph()
            continue

        # Blockquote（以 > 開頭）
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # 左側淡綠色邊框
            p_pr = p._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "24")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), "94A3B8")
            p_bdr.append(left)
            p_pr.append(p_bdr)
            joined = " ".join(quote_lines)
            add_inline(p, joined, base_size=11)
            continue

        # Checkbox list — 轉成 ☐ + 文字
        m_cb = re.match(r"^(\s*)- \[\s?\]\s+(.*)$", line)
        if m_cb:
            indent_spaces = len(m_cb.group(1))
            text = m_cb.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5 + indent_spaces * 0.3)
            p.paragraph_format.space_after = Pt(2)
            # 加一個大的 checkbox 字元
            cb_run = p.add_run("☐  ")
            apply_font(cb_run, size_pt=13)
            add_inline(p, text, base_size=11)
            i += 1
            continue

        # 已勾選 checkbox
        m_cbx = re.match(r"^(\s*)- \[x\]\s+(.*)$", line, re.IGNORECASE)
        if m_cbx:
            indent_spaces = len(m_cbx.group(1))
            text = m_cbx.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5 + indent_spaces * 0.3)
            p.paragraph_format.space_after = Pt(2)
            cb_run = p.add_run("☑  ")
            apply_font(cb_run, size_pt=13)
            add_inline(p, text, base_size=11)
            i += 1
            continue

        # Unordered list (-)
        m_ul = re.match(r"^(\s*)- (.*)$", line)
        if m_ul:
            indent_spaces = len(m_ul.group(1))
            text = m_ul.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5 + indent_spaces * 0.3)
            p.paragraph_format.space_after = Pt(2)
            bullet_run = p.add_run("• ")
            apply_font(bullet_run, size_pt=11)
            add_inline(p, text, base_size=11)
            i += 1
            continue

        # Ordered list (1. / 2. / ...)
        m_ol = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m_ol:
            indent_spaces = len(m_ol.group(1))
            num = m_ol.group(2)
            text = m_ol.group(3)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5 + indent_spaces * 0.3)
            p.paragraph_format.space_after = Pt(2)
            num_run = p.add_run(f"{num}. ")
            apply_font(num_run, size_pt=11, bold=True)
            add_inline(p, text, base_size=11)
            i += 1
            continue

        # 一般段落
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_inline(p, stripped, base_size=11)
        i += 1

    doc.save(docx_path)
    print(f"已存：{docx_path}")


if __name__ == "__main__":
    md_to_docx(
        r"C:\projects\lawyer-dashboard\scripts\briefs\林桑羽_1on1_工作表.md",
        r"C:\projects\lawyer-dashboard\scripts\briefs\林桑羽_1on1_工作表.docx",
    )
